# process_additional_documents.py
"""
Script to process additional documents in App/ (Word and Markdown files not already consolidated).
- Converts .docx and .md files to text and appends to Consolidated_Documents_with_Checklist.docx
"""
import os
from docx import Document

APP_DIR = os.path.dirname(__file__)
DOCS_DIR = os.path.join(APP_DIR, 'documents')
OUTPUT_FILE = os.path.join(DOCS_DIR, 'Consolidated_Documents_with_Checklist.docx')

# Find additional .docx and .md files
additional_files = []
for fname in os.listdir(APP_DIR):
    if fname.lower().endswith('.docx') and fname != 'documents/Consolidated_Documents_with_Checklist.docx':
        additional_files.append(os.path.join(APP_DIR, fname))
for fname in os.listdir(os.path.join(APP_DIR, 'markdown')):
    if fname.lower().endswith('.md'):
        additional_files.append(os.path.join(APP_DIR, 'markdown', fname))

def append_to_docx(docx_path, files):
    doc = Document(docx_path)
    for f in files:
        doc.add_page_break()
        doc.add_heading(os.path.basename(f), level=1)
        if f.lower().endswith('.docx'):
            d = Document(f)
            for para in d.paragraphs:
                if para.text.strip():
                    doc.add_paragraph(para.text)
        elif f.lower().endswith('.md'):
            with open(f, encoding='utf-8') as mdfile:
                for line in mdfile:
                    if line.strip():
                        doc.add_paragraph(line.strip())
    doc.save(docx_path)
    print(f'Appended {len(files)} files to {docx_path}')

if __name__ == '__main__':
    append_to_docx(OUTPUT_FILE, additional_files)
