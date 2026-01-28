# consolidate_documents.py

"""
Script to consolidate all .docx files in App/documents into a single Word document without duplication.
- Extracts text from each .docx file
- Deduplicates content
- Organizes by filename as section headers
- Outputs: App/documents/Consolidated_Documents.docx
"""

import os
from docx import Document
from collections import OrderedDict
import hashlib

DOCS_DIR = os.path.join(os.path.dirname(__file__), 'documents')
OUTPUT_FILE = os.path.join(DOCS_DIR, 'Consolidated_Documents.docx')

def get_docx_files(directory):
    return [f for f in os.listdir(directory) if f.lower().endswith('.docx')]

def hash_paragraph(text):
    return hashlib.sha256(text.strip().encode('utf-8')).hexdigest()

def extract_unique_content(files):
    seen_hashes = set()
    content_by_file = OrderedDict()
    for file in files:
        doc = Document(os.path.join(DOCS_DIR, file))
        unique_paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                h = hash_paragraph(text)
                if h not in seen_hashes:
                    seen_hashes.add(h)
                    unique_paragraphs.append(text)
        if unique_paragraphs:
            content_by_file[file] = unique_paragraphs
    return content_by_file

def create_consolidated_doc(content_by_file):
    out_doc = Document()
    out_doc.add_heading('Consolidated Documents', 0)
    for filename, paragraphs in content_by_file.items():
        out_doc.add_heading(filename, level=1)
        for para in paragraphs:
            out_doc.add_paragraph(para)
    out_doc.save(OUTPUT_FILE)
    print(f'Consolidated document created: {OUTPUT_FILE}')

def main():
    files = get_docx_files(DOCS_DIR)
    content_by_file = extract_unique_content(files)
    create_consolidated_doc(content_by_file)

if __name__ == '__main__':
    main()
