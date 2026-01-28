# consolidate_documents_with_checklist.py
"""
Script to create a new consolidated Word document with a checklist and current status summary at the top, followed by all unique content from the previous consolidation.
- Reads: App/documents/Consolidated_Documents.docx
- Writes: App/documents/Consolidated_Documents_with_Checklist.docx
"""
import os
from docx import Document

DOCS_DIR = os.path.join(os.path.dirname(__file__), 'documents')
INPUT_FILE = os.path.join(DOCS_DIR, 'Consolidated_Documents.docx')
OUTPUT_FILE = os.path.join(DOCS_DIR, 'Consolidated_Documents_with_Checklist.docx')

CHECKLIST = [
    'All .md files in root moved to organized docs/ subfolders',
    'All PowerShell scripts moved to scripts/ and references updated',
    'App folder organized: .docx, .md, .py, .jpg files in respective subfolders',
    'All .docx files in App/documents consolidated into a single deduplicated Word document',
    'No duplicate content in consolidated document',
    'Workspace structure verified and summarized',
    'All documentation and script references updated',
    'Empty folders removed',
    'Workspace is now professional and maintainable',
]

STATUS = (
    'Current State Summary:\n'
    '- Workspace root is clean, with only README.md and essential folders.\n'
    '- All documentation is organized in docs/ by category.\n'
    '- All scripts are in scripts/.\n'
    '- App/ contains only organized subfolders: documents, markdown, scripts, Images.\n'
    '- All .docx content is deduplicated and consolidated.\n'
    '- No files lost; all references are up to date.'
)

def main():
    # Read the original consolidated document
    doc = Document(INPUT_FILE)
    # Create new document
    out_doc = Document()
    out_doc.add_heading('Vets Ready Consolidated Documentation', 0)
    out_doc.add_heading('Project Checklist', level=1)
    for item in CHECKLIST:
        out_doc.add_paragraph(item, style='List Bullet')
    out_doc.add_heading('Current State Summary', level=1)
    for line in STATUS.split('\n'):
        out_doc.add_paragraph(line)
    out_doc.add_page_break()
    # Copy all content from the original consolidated document (skip its title)
    skip_title = True
    for para in doc.paragraphs:
        if skip_title and para.style.name.startswith('Heading') and para.text.strip():
            skip_title = False
            continue
        if not skip_title:
            out_doc.add_paragraph(para.text, style=para.style)
    out_doc.save(OUTPUT_FILE)
    print(f'Checklist and summary document created: {OUTPUT_FILE}')

if __name__ == '__main__':
    main()
