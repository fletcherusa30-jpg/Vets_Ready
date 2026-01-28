from docx import Document
from docx.shared import Pt

# Paths
consolidated_path = r"c:/Dev/Vets Ready/App/documents/Consolidated_Documents_with_Checklist.docx"
addendum_title = "VetsReady Engineering Addendum (Appended January 27, 2026)"
addendum_content = [
    "1. Engineering Philosophy",
    "- Zero‑Placeholder Rule",
    "- Modular, Composable Architecture",
    "- Shared Types Across Frontend + Backend",
    "- Predictable Folder Structure",
    "2. Resume Builder — Engineering Addendum",
    "2.1 Required Data Models",
    "- Resume",
    "- ExperienceItem",
    "- MOSMapping",
    "2.2 Required Backend Endpoints",
    "- POST /resume/generate",
    "- POST /resume/tailor",
    "- POST /resume/achievements",
    "- GET /mos/translate/{mosCode}",
    "2.3 Required Frontend Components",
    "- ResumeEditor.tsx",
    "- ResumePreview.tsx",
    "- AchievementGenerator.tsx",
    "- JobTargetingPanel.tsx",
    "2.4 Required AI Pipelines",
    "- Resume Generation",
    "- Job Targeting",
    "3. Job Recruiting Platform — Engineering Addendum",
    "3.1 Required Data Models",
    "- Job",
    "- VeteranProfile",
    "3.2 Required Backend Endpoints",
    "- POST /jobs/search",
    "- POST /employers/search",
    "- GET /certifications/pathway/{certName}",
    "3.3 Required Frontend Components",
    "- MOSSelector.tsx",
    "- JobListingCard.tsx",
    "- VeteranProfileCard.tsx",
    "4. Budget & Retirement Tools — Engineering Addendum",
    "4.1 Required Data Models",
    "- Budget",
    "- RetirementInputs",
    "4.2 Required Backend Endpoints",
    "- POST /budget/calculate",
    "- POST /retirement/calculate",
    "4.3 Required Frontend Components",
    "- BudgetPage.tsx",
    "- RetirementPage.tsx",
    "5. PowerShell Automation — Engineering Addendum",
    "5.1 Required Scripts",
    "- bootstrap.ps1",
    "- integrity-check.ps1",
    "- rebuild.ps1",
    "6. Copilot Behavioral Addendum",
    "- Always generate real code",
    "- Always wire UI to logic",
    "- Always enforce modularity",
    "- Always generate tests",
    "7. Final Summary",
    "This Engineering Addendum defines the deep technical requirements for all VetsReady modules. It ensures Copilot and developers produce real, functional code, modular architecture, shared data models, fully wired UI, AI‑powered pipelines, and professional‑grade tools.",
    "This is the engineering backbone of the entire VetsReady ecosystem."
]

def append_addendum():
    doc = Document(consolidated_path)
    doc.add_page_break()
    doc.add_heading(addendum_title, level=1)
    for line in addendum_content:
        if line.strip().endswith(":"):
            doc.add_heading(line, level=2)
        elif line.startswith("-"):
            doc.add_paragraph(line, style='List Bullet')
        else:
            doc.add_paragraph(line)
    doc.save(consolidated_path)
    print("Addendum appended successfully.")

if __name__ == "__main__":
    append_addendum()
