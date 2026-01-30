"""
DOCUMENT PROCESSING UTILITY

Extracts text from Word documents (.docx) and processes them
through the document scanner system.

FEATURES:
- Word document text extraction
- Metadata extraction (author, created date, modified date)
- Table content extraction
- Header/footer extraction
- Save results to structured format

USAGE:
    python process_word_docs.py --file "path/to/document.docx"
    python process_word_docs.py --scan-all
"""

import sys
import os
import json
import argparse
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List, Optional
import logging

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    logger.error("python-docx not installed. Install with: pip install python-docx")
    DOCX_AVAILABLE = False


class WordDocumentProcessor:
    """Process Word documents and extract content"""

    def __init__(self, output_dir: str = "./data/extracted"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract all text content from Word document

        Returns:
            {
                'success': bool,
                'file_path': str,
                'text': str,
                'paragraphs': List[str],
                'tables': List[Dict],
                'metadata': Dict,
                'character_count': int,
                'word_count': int,
                'paragraph_count': int,
                'table_count': int
            }
        """
        if not DOCX_AVAILABLE:
            return {
                'success': False,
                'error': 'python-docx library not available'
            }

        file_path = Path(file_path)

        if not file_path.exists():
            return {
                'success': False,
                'error': f'File not found: {file_path}'
            }

        logger.info(f"Processing Word document: {file_path}")

        try:
            doc = Document(file_path)

            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Only add non-empty paragraphs
                    paragraphs.append(text)

            # Extract tables
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = {
                    'index': table_idx,
                    'rows': [],
                    'row_count': len(table.rows),
                    'col_count': len(table.columns) if table.rows else 0
                }

                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data['rows'].append(row_data)

                tables.append(table_data)

            # Combine all text
            full_text = '\n\n'.join(paragraphs)

            # Add table text
            for table in tables:
                table_text = '\n'.join([' | '.join(row) for row in table['rows']])
                full_text += f"\n\n[TABLE {table['index']}]\n{table_text}"

            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                'author': core_props.author or 'Unknown',
                'title': core_props.title or file_path.stem,
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'created': str(core_props.created) if core_props.created else None,
                'modified': str(core_props.modified) if core_props.modified else None,
                'last_modified_by': core_props.last_modified_by or '',
                'revision': core_props.revision or 0
            }

            # Statistics
            word_count = len(full_text.split())
            character_count = len(full_text)

            result = {
                'success': True,
                'file_path': str(file_path),
                'file_name': file_path.name,
                'text': full_text,
                'paragraphs': paragraphs,
                'tables': tables,
                'metadata': metadata,
                'character_count': character_count,
                'word_count': word_count,
                'paragraph_count': len(paragraphs),
                'table_count': len(tables),
                'extracted_at': datetime.utcnow().isoformat()
            }

            logger.info(f"Extracted {character_count} characters, {word_count} words, "
                       f"{len(paragraphs)} paragraphs, {len(tables)} tables")

            return result

        except Exception as e:
            logger.error(f"Failed to process document: {e}")
            return {
                'success': False,
                'error': str(e),
                'file_path': str(file_path)
            }

    def save_extraction(self, extraction_result: Dict[str, Any], format: str = 'json') -> str:
        """
        Save extraction results to file

        Args:
            extraction_result: Result from extract_text()
            format: Output format ('json', 'txt', 'both')

        Returns:
            Path to saved file(s)
        """
        if not extraction_result.get('success'):
            logger.error("Cannot save failed extraction")
            return None

        file_name = extraction_result['file_name']
        base_name = Path(file_name).stem
        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')

        saved_files = []

        if format in ['json', 'both']:
            # Save full JSON
            json_path = self.output_dir / f"{base_name}_extracted_{timestamp}.json"
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(extraction_result, f, indent=2, ensure_ascii=False)
            logger.info(f"Saved JSON to: {json_path}")
            saved_files.append(str(json_path))

        if format in ['txt', 'both']:
            # Save plain text
            txt_path = self.output_dir / f"{base_name}_extracted_{timestamp}.txt"
            with open(txt_path, 'w', encoding='utf-8') as f:
                # Write metadata header
                f.write("=" * 80 + "\n")
                f.write(f"Document: {extraction_result['file_name']}\n")
                f.write(f"Extracted: {extraction_result['extracted_at']}\n")
                f.write(f"Words: {extraction_result['word_count']}\n")
                f.write(f"Characters: {extraction_result['character_count']}\n")
                f.write("=" * 80 + "\n\n")

                # Write content
                f.write(extraction_result['text'])

            logger.info(f"Saved text to: {txt_path}")
            saved_files.append(str(txt_path))

        return saved_files

    def scan_directory(self, directory: str, pattern: str = "*.docx") -> List[str]:
        """
        Scan directory for Word documents

        Returns:
            List of document paths
        """
        directory = Path(directory)

        # Exclude patterns
        exclude_patterns = [
            'node_modules',
            '.venv',
            '_archive',
            '.git',
            '__pycache__',
            'venv'
        ]

        documents = []
        for file_path in directory.rglob(pattern):
            # Check if path contains any exclude pattern
            if any(exclude in str(file_path) for exclude in exclude_patterns):
                continue
            documents.append(str(file_path))

        return documents


def main():
    """Main execution"""
    parser = argparse.ArgumentParser(description='Process Word documents')
    parser.add_argument('--file', type=str, help='Path to specific document')
    parser.add_argument('--scan-all', action='store_true', help='Scan entire workspace')
    parser.add_argument('--directory', type=str, default='.', help='Directory to scan')
    parser.add_argument('--output', type=str, default='./data/extracted', help='Output directory')
    parser.add_argument('--format', type=str, choices=['json', 'txt', 'both'], default='both',
                       help='Output format')

    args = parser.parse_args()

    processor = WordDocumentProcessor(output_dir=args.output)

    if not DOCX_AVAILABLE:
        print("❌ python-docx not installed")
        print("Install with: pip install python-docx")
        sys.exit(1)

    documents_to_process = []

    if args.file:
        # Process single file
        documents_to_process.append(args.file)
    elif args.scan_all:
        # Scan workspace
        print(f"🔍 Scanning {args.directory} for Word documents...")
        documents_to_process = processor.scan_directory(args.directory)
        print(f"Found {len(documents_to_process)} documents")
    else:
        parser.print_help()
        sys.exit(0)

    if not documents_to_process:
        print("No documents to process")
        sys.exit(0)

    # Process documents
    results = []
    success_count = 0
    fail_count = 0

    print(f"\n🔄 Processing {len(documents_to_process)} documents...\n")

    for doc_path in documents_to_process:
        print(f"📄 Processing: {Path(doc_path).name}")

        # Extract text
        result = processor.extract_text(doc_path)

        if result['success']:
            # Save extraction
            saved_files = processor.save_extraction(result, format=args.format)
            result['saved_files'] = saved_files

            print(f"  ✅ Success - {result['character_count']} chars, "
                  f"{result['word_count']} words")
            print(f"  📁 Saved to: {', '.join(saved_files)}")
            success_count += 1
        else:
            print(f"  ❌ Failed: {result.get('error', 'Unknown error')}")
            fail_count += 1

        results.append(result)
        print()

    # Summary
    print("=" * 80)
    print("📊 PROCESSING SUMMARY")
    print("=" * 80)
    print(f"Total Documents: {len(documents_to_process)}")
    print(f"✅ Successful: {success_count}")
    print(f"❌ Failed: {fail_count}")

    # Save summary
    summary_path = processor.output_dir / f"summary_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.json"
    with open(summary_path, 'w', encoding='utf-8') as f:
        json.dump({
            'timestamp': datetime.utcnow().isoformat(),
            'total': len(documents_to_process),
            'success': success_count,
            'failed': fail_count,
            'results': results
        }, f, indent=2, ensure_ascii=False)

    print(f"\n📋 Summary saved to: {summary_path}")
    print("\n✅ Processing complete!\n")


if __name__ == '__main__':
    main()
